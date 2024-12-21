import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
st.set_page_config(layout="wide")

# Function to reset session state
def reset_app():
    for key in list(st.session_state.keys()):
        del st.session_state[key]

# Check if the reset button was clicked
if st.button("Restart App"):
    reset_app()

df = pd.read_csv("CompanyDetails.csv")
df = df.dropna()
ExistingCompany = df["Company Name"].to_list()
ExistingCompanyAdd = df["Address"].to_list()

# Title bar
st.markdown(
        """<h3 style='text-align: center; color: white; background-color: maroon; padding: 10px;'>Tilak Raj Tour & Travel</h3>
        <h6 style='text-align: center; color: white; background-color: maroon; padding: 10px;'>Add:Noida,U.P. 201301, Contact: 9711981687,8368200227</h3>""",
        unsafe_allow_html=True
    )
st.write("")

# Function to create the Word document
def create_bill_with_dataframe(dataframe, Company, We):
    doc = Document()
    Heading = "Tilak Raj Tour & Travel\n Add: Noida, U.P., 201301,       Contact: 9711981687, 8368200227"
    # Add title
    title = doc.add_paragraph(Heading)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.runs[0].font.size = Pt(14)
    title.runs[0].bold = True

    # Add "Bill To" and "Bill From" section
    table = doc.add_table(rows=1, cols=2)
    #table.style = 'Table Grid'

    # Add "Bill To"
    cell1 = table.cell(0, 0)
    cell1.text = f"Bill To:\n{Company}"
    cell1.paragraphs[0].runs[0].bold = True
    cell1.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Add "Bill From"
    cell2 = table.cell(0, 1)
    cell2.text = f"Bill From:\n{We}"
    cell2.paragraphs[0].runs[0].bold = True
    cell2.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Add a space
    doc.add_paragraph()
    note_paragraph = doc.add_paragraph("Quotations are Mentioned Below:")
    note_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    note_paragraph.runs[0].font.size = Pt(10)
    note_paragraph.runs[0].italic = True
    # Add a table for the Pandas DataFrame
    df_table = doc.add_table(rows=1, cols=len(dataframe.columns))
    df_table.style = 'Table Grid'

    # Add column headers
    hdr_cells = df_table.rows[0].cells
    for i, column_name in enumerate(dataframe.columns):
        hdr_cells[i].text = str(column_name)
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add data rows
    for _, row in dataframe.iterrows():
        row_cells = df_table.add_row().cells
        for i, cell_value in enumerate(row):
            row_cells[i].text = str(cell_value)
            row_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add a note at the bottom
    doc.add_paragraph("\n")
    note_paragraph = doc.add_paragraph("Note: All Charges are exclusive of MCD or any other types Toll Taxes.")
    note_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    note_paragraph.runs[0].font.size = Pt(10)
    note_paragraph.runs[0].italic = True

    
    # Save document to a BytesIO buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Company Dropdown
CompanyName = st.selectbox("Select Company : कंपनी चुनें",ExistingCompany, key="dropdown1")
index = ExistingCompany.index(CompanyName)
try:
    CompanyAddress = ExistingCompanyAdd[index]
except:
    CompanyAddress = "N/A"
col1,col2,col3 = st.columns([1,1,1])
# Initialize session state to store the list of items
if 'items' not in st.session_state:
    st.session_state['items'] = []
with col1:
    # Create the form for item entry
    with st.form(key='item_form'):
        # Route Details
        Route_Details = st.text_input("Route Details : यात्रा विवरण")
    
        # CAB Dropdown
        Cab_Dropdown = st.radio("Select car : कार चुनें",("Ertiga : अर्टिगा, 7 seater", "Eeco :ईको, 7 seater", "Dzire :डिजायर  or अन्य 5 seater", "All: सभी"),)
        
        TempDict = {}
        if "Ertiga : अर्टिगा, 7 seater" == Cab_Dropdown:
            Ertiga_amount = st.number_input("Amount for Ertiga: अर्टिगा के लिए राशि")
            Dzire_amount = None
            Eeco_amount = None
        elif "Eeco :ईको, 7 seater" == Cab_Dropdown:
            Eeco_amount = st.number_input("Amount for Eeco: इको के लिए राशि")
            Dzire_amount = None
            Ertiga_amount = None
        elif "Dzire :डिजायर  or अन्य 5 seater" == Cab_Dropdown:
            Dzire_amount =  st.number_input("Amount for Dzire: डिजायर के लिए रकम")
            Eeco_amount = None
            Ertiga_amount = None
        else:
            Ertiga_amount = st.number_input("Amount for Ertiga: अर्टिगा के लिए राशि")
            Eeco_amount = st.number_input("Amount for Eeco: इको के लिए राशि")
            Dzire_amount = st.number_input("Amount for Dzire: डिजायर के लिए रकम")
        
        try:
            TempDict["Cab"] = Cab_Dropdown
            if Ertiga_amount > 0:
                TempDict["Amount"] = Cab_Dropdown
            elif Eeco_amount > 0:
                TempDict["Amount"] = Cab_Dropdown
            elif Dzire_amount > 0:
                TempDict["Amount"] = Cab_Dropdown
            else:
                TempDict["Amount"] = "NA"
        except:
            TempDict["Amount"] = "NA"
        check = set(list(TempDict.values()))
        if len(check)==1:
            # Submit button for adding an item
            submit_button = st.form_submit_button(label='Add New Cab : नई कैब जोड़ें')
        else:
            # Submit button for adding an item
            submit_button = st.form_submit_button(label='Add Amount : राशि जोड़ें')
        TotalLen = len(st.session_state['items'])
        # When the user submits the form to add the item, append the item to the list
        if submit_button:
            if Route_Details and Cab_Dropdown:
                # Add item to the session state list
                st.session_state['items'].append({'Route Details': Route_Details,'Selected cab': Cab_Dropdown,'Ertiga Amount' : Ertiga_amount,'Eeco Amount' : Eeco_amount,'Dzire Amount' : Dzire_amount })
                # Reset the form fields by simply not using st.experimental_rerun
                # The next cycle will show the updated state without rerunning the script manually
                
            else:
                st.error("Please enter Route item details: अमान्य मार्ग")
        
    corruptIndex = []
    # Display all added items
    if st.session_state['items']:
        total_amount = 0
        st.subheader("Added Items")
        Success = False
        itemnumber = 1
        for i, item in enumerate(st.session_state['items']):
            lst = [item['Ertiga Amount'], item['Eeco Amount'], item['Dzire Amount']]
            Filteredlst = [x for x in lst if x is not None]
            if len(Filteredlst)<len(lst) and sum(Filteredlst)>0:
                st.write(f"{itemnumber}. - **Route Details**: {item['Route Details']}" )
                st.write(f"- **Selected cab**: {item['Selected cab']}  ")
                if item['Ertiga Amount']:
                    st.write(f"- **Ertiga Amount**: {item['Ertiga Amount']} ")
                if item['Eeco Amount']:
                    st.write(f"- **Eeco Amount**: {item['Eeco Amount']}" )
                if item['Dzire Amount']:
                    st.write(f"- **Dzire Amount**: {item['Dzire Amount']}")
                Success = True
                itemnumber = itemnumber+1
            elif len(Filteredlst)==len(lst):
                if min(lst)!=0.0:
                    Success = True
                    st.write(f"{itemnumber}. - **Route Details**: {item['Route Details']}" )
                    st.write(f"- **Selected cab**: {item['Selected cab']}  ")
                    if item['Ertiga Amount']:
                        st.write(f"- **Ertiga Amount**: {item['Ertiga Amount']} ")
                    if item['Eeco Amount']:
                        st.write(f"- **Eeco Amount**: {item['Eeco Amount']}" )
                    if item['Dzire Amount']:
                        st.write(f"- **Dzire Amount**: {item['Dzire Amount']}")
                    itemnumber = itemnumber+1
                    st.warning("Next Entry will be corrupt but don't worry, It;ll be handle in Final quotation (अगली प्रविष्टि भ्रष्ट हो जाएगी लेकिन चिंता न करें, इसे अंतिम उद्धरण में संभाल लिया जाएगा)")
                else:
                    Success = False
                    corruptIndex.append(i)
            else:
                Success = False
                corruptIndex.append(i)
        if Success:
            st.write("Route added! You can now add another Route:  मार्ग जोड़ा गया! अब आप दूसरा रूट जोड़ सकते हैं.")


# Generate Bill Button (to finalize the bill)
if len(st.session_state['items']) > 0:
    
    with col2:
        generate_bill_button = st.button("Generate Quotation")
        if generate_bill_button: 
            # Display the final bill
            st.header("Final Quotation")
            st.write("")
            st.subheader("**Bill from:**")
            st.write("**Tilak Raj Tour and Travel**")
            st.write("Contact:      9711981687, 8368200227")
            st.write("")
        with col3:
            st.header("")
            st.write("")
            st.subheader("**Bill for:**")  
            st.write(f"**{CompanyName}**")
            st.write(f"Address:     {CompanyAddress}")
            st.write("")
        with col2:
            st.subheader("Quotation Details")
            RouteList = []
            Ertiga = []
            Eeco = []
            Dzire = []
            # Print out the item list
            for item in st.session_state['items']:
                try:
                    RouteList.append(item['Route Details'])
                    Ertiga.append(item['Ertiga Amount'])
                    Dzire.append(item['Dzire Amount'])
                    Eeco.append(item['Eeco Amount'])
                except:
                    RouteList.append(None)
                    Ertiga.append(None)
                    Dzire.append(None)
                    Eeco.append(None)  
                    
            #Handling Corrupt Entries
            DetailsDataframe = pd.DataFrame({"Route Name":RouteList, "Ertiga":Ertiga ,"Eeco" :Eeco, "Dzire or other" : Dzire})
            DetailsDataframe.drop(corruptIndex, axis=0, inplace=True)
            DetailsDataframe = DetailsDataframe.reset_index(drop=True)
            
            try:
                #Hanling next all entries
                new_corrupt = DetailsDataframe.index[DetailsDataframe.notnull().all(axis=1)].tolist()
                new_corrupt = [i+1 for i in new_corrupt]
                DetailsDataframe.drop(new_corrupt, axis=0, inplace=True)
                DetailsDataframe = DetailsDataframe.reset_index(drop=True)
            except Exception as e:
                st.write("Next Entry will be corrupt but don't worry, It;ll be handle in Final quotation (अगली प्रविष्टि भ्रष्ट हो जाएगी लेकिन चिंता न करें, इसे अंतिम उद्धरण में संभाल लिया जाएगा)")
            st.write(DetailsDataframe)
            
        with col3:
            # Button to generate and download the Word document
            Company = CompanyName + "\n" + CompanyAddress
            We = "Tilak Raj Tour and Travel" + "\n" + "Add: Noida Sec 49, U.P., 201301"
            buffer = create_bill_with_dataframe(DetailsDataframe, Company, We)
            tempname = CompanyName + "_Quotation.docx"
            st.download_button(
                    label="Download Quotation",
                    data=buffer,
                    file_name=tempname,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            
        
